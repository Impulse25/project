"""
edu/includes/vedomost_builder.py
Генерирует зачётную/экзаменационную ведомость на основе шаблона.

Вызов:
  python vedomost_builder.py --json data.json --out output.docx
"""
import json, sys, argparse, copy, shutil, os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE = os.path.join(os.path.dirname(__file__), '..', 'templates', 'vedomost_template.docx')


def set_para_text(para, new_text):
    """Заменить текст параграфа, сохраняя форматирование первого run."""
    if not para.runs:
        para.text = new_text
        return
    # Сохраняем форматирование первого run
    first_run = para.runs[0]
    rpr = first_run._r.find(qn('w:rPr'))
    # Очищаем все runs
    for run in para.runs:
        run.text = ''
    # Пишем текст в первый run
    first_run.text = new_text


def replace_in_para(para, old, new):
    """Заменить old→new в параграфе с учётом разбивки по run-ам."""
    full = para.text
    if old not in full:
        return False
    new_full = full.replace(old, new)
    set_para_text(para, new_full)
    return True


def clear_table_row(row):
    """Очищает текст всех ячеек строки."""
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.text = ''


def copy_row_format(src_row, dst_row):
    """Копирует XML форматирования строки (высоту, свойства ячеек)."""
    src_trPr = src_row._tr.find(qn('w:trPr'))
    if src_trPr is not None:
        dst_trPr = dst_row._tr.find(qn('w:trPr'))
        if dst_trPr is not None:
            dst_row._tr.remove(dst_trPr)
        dst_row._tr.insert(0, copy.deepcopy(src_trPr))


def add_row_like(table, template_row):
    """Добавляет новую строку в таблицу, скопировав структуру template_row."""
    new_tr = copy.deepcopy(template_row._tr)
    table._tbl.append(new_tr)
    return table.rows[-1]


def fill_cell(cell, text, bold=None):
    """Заполнить ячейку текстом, сохраняя форматирование."""
    for para in cell.paragraphs:
        if para.runs:
            para.runs[0].text = str(text) if text else ''
            for run in para.runs[1:]:
                run.text = ''
            if bold is not None:
                para.runs[0].font.bold = bold
        else:
            para.text = str(text) if text else ''
        break


def build(data: dict, out_path: str):
    template_path = os.path.abspath(TEMPLATE)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f'Шаблон не найден: {template_path}')

    doc = Document(template_path)

    discipline   = data.get('discipline', '')
    group_name   = data.get('group_name', '')
    specialty    = data.get('specialty', '')
    qualification= data.get('qualification', '')
    teacher      = data.get('teacher', '')
    course_num   = data.get('course_num', '')
    exam_type    = data.get('exam_type', 'зачёт')
    counts       = data.get('counts', {}) or {}
    students     = data.get('students', [])

    # ── Заменяем текст в параграфах ──────────────────────────────────────────
    for para in doc.paragraphs:
        txt = para.text

        if 'ЗАЧЕТНАЯ ВЕДОМОСТЬ' in txt or 'ЗАЧЁТНАЯ ВЕДОМОСТЬ' in txt or 'ЭКЗАМЕНАЦИОННАЯ ВЕДОМОСТЬ' in txt:
            set_para_text(para, 'ЭКЗАМЕНАЦИОННАЯ ВЕДОМОСТЬ' if exam_type == 'экзамен' else 'ЗАЧЕТНАЯ ВЕДОМОСТЬ')
            continue

        # Индекс дисциплины
        if 'Индекс модуля' in txt and 'модулю' in txt:
            # Находим часть после "модулю "
            prefix = 'Индекс модуля, по дисциплине и (или) модулю '
            set_para_text(para, prefix + discipline)
            continue

        # Курс и группа
        if 'курса группы' in txt:
            set_para_text(para, f'«{course_num}» курса группы {group_name}')
            continue

        # Специальность
        if txt.startswith('Специальность:'):
            set_para_text(para, f'Специальность: {specialty}')
            continue

        # Квалификация
        if txt.startswith('Квалификация:'):
            set_para_text(para, f'Квалификация: {qualification}')
            continue

        # Преподаватель
        if txt.startswith('Преподаватель:'):
            set_para_text(para, f'Преподаватель: {teacher}')
            continue

    # ── Таблица студентов ────────────────────────────────────────────────────
    table = doc.tables[0]
    # Строки 0,1,2 — заголовки. Строки 3+ — студенты
    header_rows = 3
    data_rows_in_template = len(table.rows) - header_rows  # обычно 19

    # Сохраняем эталонную строку-шаблон (строка 3) для добавления новых
    template_data_row = table.rows[header_rows]

    n_students = len(students)

    # Расширяем или сокращаем таблицу до нужного количества строк
    current_data_rows = len(table.rows) - header_rows

    if n_students > current_data_rows:
        # Добавляем строки
        for _ in range(n_students - current_data_rows):
            add_row_like(table, template_data_row)
    elif n_students < current_data_rows:
        # Удаляем лишние строки снизу
        for _ in range(current_data_rows - n_students):
            row_to_del = table.rows[-1]._tr
            table._tbl.remove(row_to_del)

    # Заполняем строки студентов
    for si, student in enumerate(students):
        row = table.rows[header_rows + si]
        # Очищаем строку
        clear_table_row(row)
        # Заполняем
        fill_cell(row.cells[0],  si + 1)
        if len(row.cells) >= 15:
            fill_cell(row.cells[1],  student.get('rating_score', ''))
            fill_cell(row.cells[2],  student.get('rating_letter', ''))
            fill_cell(row.cells[3],  student.get('rating_gpa', ''))
            fill_cell(row.cells[4],  student.get('full_name', ''))
            fill_cell(row.cells[5],  student.get('written_score', ''))
            fill_cell(row.cells[6],  student.get('written_letter', ''))
            fill_cell(row.cells[7],  student.get('written_gpa', ''))
            fill_cell(row.cells[8],  student.get('oral_score', ''))
            fill_cell(row.cells[9],  student.get('oral_letter', ''))
            fill_cell(row.cells[10], student.get('oral_gpa', ''))
            fill_cell(row.cells[11], student.get('total_score', ''))
            fill_cell(row.cells[12], student.get('total_letter', ''))
            fill_cell(row.cells[13], student.get('total_gpa', ''))
            fill_cell(row.cells[14], '')
        else:
            fill_cell(row.cells[1],  student.get('rating_letter', ''))
            fill_cell(row.cells[2],  student.get('rating_gpa', ''))
            fill_cell(row.cells[3],  student.get('ticket_num', ''))
            fill_cell(row.cells[4],  student.get('full_name', ''))
            fill_cell(row.cells[5],  student.get('written_letter', ''))
            fill_cell(row.cells[6],  student.get('written_gpa', ''))
            fill_cell(row.cells[7],  student.get('oral_letter', ''))
            fill_cell(row.cells[8],  student.get('oral_gpa', ''))
            fill_cell(row.cells[9],  student.get('total_gpa', ''))
            fill_cell(row.cells[10], student.get('total_letter', ''))
            fill_cell(row.cells[11], student.get('total_gpa2', ''))
            fill_cell(row.cells[12], '')  # подпись

    # Обновляем блок с количеством оценок, если он есть в шаблоне.
    if len(doc.tables) > 1:
        for table_footer in doc.tables[1:]:
            for row in table_footer.rows:
                for cell in row.cells:
                    if 'Количество оценок' in cell.text:
                        cell.text = (
                            'Количество оценок:\n'
                            f'А, А- ___{int(counts.get("excellent", 0) or 0)}___\n'
                            f'В+, В, В-, С+ ___{int(counts.get("good", 0) or 0)}___\n'
                            f'С, С- D+, D ___{int(counts.get("satisfactory", 0) or 0)}___\n'
                            f'F ___{int(counts.get("fail", 0) or 0)}___'
                        )

    doc.save(out_path)
    print(f'OK:{out_path}')


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('--json', required=True)
    parser.add_argument('--out',  required=True)
    args = parser.parse_args()
    with open(args.json, encoding='utf-8-sig') as f:
        data = json.load(f)
    build(data, args.out)
